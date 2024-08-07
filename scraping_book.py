import requests
from docx import Document
from docx.shared import Inches
import os
import shutil
import json
from PIL import Image
from io import BytesIO

tipo_persona = [{"alum4" : "23EPMAT4"}]

tipo_libro = [{"RM" : "A9"}]

# Directorio
base_dir = "D:\\Usuario\\Downloads"
temp_dir = os.path.join(base_dir, "temp_images")

historial = []

# Limpiar el directorio temporal al inicio
if os.path.exists(temp_dir):
    shutil.rmtree(temp_dir)
os.makedirs(temp_dir, exist_ok=True)

for i, diccionario_persona in enumerate(tipo_persona):
    for clave_persona, valor_persona in diccionario_persona.items():
        for diccionario_libro in tipo_libro:
            for clave_libro, valor_libro in diccionario_libro.items():
                print(valor_persona, valor_libro)
                doc = Document()
                for x in range(1, 21):
                    print(f'x: {x}')
                    for y in range(1, 11):
                        #print(f'y: {y}')
                        consecutive_misses = 0
                        for z in range(1, 201):
                            if consecutive_misses >= 3:
                                break
                            url = f"https://libros.edicioneslexicom.pe/LIBROS/{valor_persona}/{valor_libro}_U{x}_T{y}_1/files/mobile/{z}.jpg?"
                            response = requests.get(url)
                            #print(response, url)
                            historial.append({
                                'valor_persona': valor_persona,
                                'valor_libro': valor_libro,
                                'x': x,
                                'y': y,
                                'z': z,
                                'status_code': response.status_code,
                                'url': url
                            })
                            
                            if response.status_code != 200:
                                consecutive_misses += 1
                                #print("error")
                                continue
                            consecutive_misses = 0
                            
                            try:
                                # Verifica si el contenido descargado es una imagen
                                image = Image.open(BytesIO(response.content))
                                image_path = os.path.join(temp_dir, f'image_{i}_{x}_{y}_{z}.jpg')
                                image.save(image_path)
                                
                                # Agregar la imagen al documento de Word
                                doc.add_picture(image_path, width=Inches(6))  # Ajusta el ancho de la imagen según sea necesario
                                doc.add_paragraph('')  # Añade un espacio entre imágenes
                            except Exception as e:
                                error_message = f"Error al procesar la imagen {url}: {e}"
                                print(error_message)
                                doc.add_paragraph(error_message)
                                continue
                                
                # Guardar el documento de Word
                doc_path = os.path.join(base_dir, f'documento_con_imagenes_{valor_persona}_{valor_libro}.docx')
                doc.save(doc_path)
                print(f"Creado con éxito:f'documento_con_imagenes_{valor_persona}_{valor_libro}")

                # Limpiar el directorio temporal después de cada iteración
                try:
                    shutil.rmtree(temp_dir)
                    os.makedirs(temp_dir, exist_ok=True)
                except FileNotFoundError:
                    print("El directorio temporal no existe o ya ha sido eliminado.")
                
# Limpiar el directorio temporal al final
try:
    shutil.rmtree(temp_dir)
except FileNotFoundError:
    print("El directorio temporal no existe o ya ha sido eliminado.")

print("DONE.")